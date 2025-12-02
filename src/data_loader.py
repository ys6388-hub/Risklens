import os
import glob
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
from moviepy.editor import VideoFileClip
from io import BytesIO

load_dotenv()


class StandardDataLoader:
    def __init__(self):
        """Initialize loader + Whisper audio if API key exists."""
        self.has_audio_capability = False
        api_key = os.getenv("OPENAI_API_KEY")

        if api_key:
            import openai
            openai.api_key = api_key
            self.has_audio_capability = True

    # ===================================================================
    # 1. LOCAL FOLDER LOADING
    # ===================================================================
    def load_dataset(self, folder_path, limit=None):
        """Scan folder and read PDF, DOCX, TXT, VIDEO, AUDIO"""
        results = []

        file_patterns = {
            "pdf": ["*.pdf"],
            "docx": ["*.docx"],
            "txt": ["*.txt"],
            "video": ["*.mp4", "*.mov", "*.avi"],
            "audio": ["*.mp3", "*.wav", "*.m4a"],
        }

        all_files = []
        for _, patterns in file_patterns.items():
            for pattern in patterns:
                all_files.extend(glob.glob(os.path.join(folder_path, "**", pattern), recursive=True))

        print(f"📁 Scanning {folder_path}... Found {len(all_files)} files")

        for fp in (all_files if limit is None else all_files[:limit]):
            try:
                processed = self._process_file(fp)
                if processed:
                    results.append(processed)
            except Exception as e:
                print(f"Skipping {fp}: {e}")

        return results

    # ===================================================================
    # 2. UPLOADED FILE (BYTES)
    # ===================================================================
    def load_pdf_from_bytes(self, file_bytes):
        try:
            pdf = PdfReader(BytesIO(file_bytes))
            return "".join([page.extract_text() or "" for page in pdf.pages])
        except Exception as e:
            print("PDF parse failed:", e)
            return ""

    def load_docx_from_bytes(self, file_bytes):
        try:
            doc = Document(BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            print("DOCX parse failed:", e)
            return ""

    def load_txt_from_bytes(self, file_bytes):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""

    # ===================================================================
    # 3. INTERNAL PROCESSOR FOR LOCAL FILES
    # ===================================================================
    def _process_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        content = ""

        if ext == ".pdf":
            content = self._read_pdf(file_path)
        elif ext == ".docx":
            content = self._read_docx(file_path)
        elif ext == ".txt":
            content = self._read_txt(file_path)
        elif ext in [".mp4", ".mov", ".avi"] and self.has_audio_capability:
            content = self._process_video(file_path)
        elif ext in [".mp3", ".wav", ".m4a"] and self.has_audio_capability:
            content = self._transcribe(file_path)

        if content and len(content.strip()) > 10:
            return {"filename": filename, "type": ext, "content": content}

        return None

    # ===================================================================
    # 4. LOCAL READERS
    # ===================================================================
    def _read_pdf(self, path):
        try:
            pdf = PdfReader(path)
            return "".join([page.extract_text() or "" for page in pdf.pages[:5]])
        except Exception as e:
            print("PDF read failed:", e)
            return ""

    def _read_docx(self, path):
        try:
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            print("DOCX read failed:", e)
            return ""

    def _read_txt(self, path):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            print("TXT read failed:", e)
            return ""

    # ===================================================================
    # 5. VIDEO HANDLING
    # ===================================================================
    def _process_video(self, path):
        try:
            clip = VideoFileClip(path)
            temp_audio = "temp_audio_buffer.mp3"
            clip.audio.write_audiofile(temp_audio, logger=None)
            text = self._transcribe(temp_audio)

            if os.path.exists(temp_audio):
                os.remove(temp_audio)

            return text
        except Exception as e:
            print("Video processing failed:", e)
            return ""

    # ===================================================================
    # 6. WHISPER TRANSCRIPTION
    # ===================================================================
    def _transcribe(self, audio_path):
        try:
            import openai
            with open(audio_path, "rb") as f:
                result = openai.Audio.transcribe("whisper-1", f)
                return result.get("text", "")
        except Exception as e:
            print("Whisper failed:", e)
            return ""
